import sys
import math
from pandas.api.types import is_numeric_dtype
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go
import numpy as np
import plotly.express as px
import pandas as pd
from plotly.subplots import make_subplots



class tableAnalysis:

    def __init__(self, cusTable, settings):

        # Getting The Title for This Section
        allKeys = list(cusTable.keys())
        if "title" in allKeys:
            self.title = cusTable['title']
        else:
            self.title = settings['tableAnalysis']['title']

        # Use Previous Results for This Section
        if "usePrevResults" in allKeys:
            self.usePrevResults = cusTable['usePrevResults']
        else:
            self.usePrevResults = settings['tableAnalysis']['usePrevResults']

        # Rescale The Table for better Analysis
        if "rescale" in allKeys:
            self.rescale = cusTable['rescale']
        else:
            self.rescale = settings['tableAnalysis']['rescale']

        # Excel File Selection
        if "excel" in allKeys:
            self.excel = cusTable['excel']
        else:
            self.excel = settings['tableAnalysis']['excel']

        # SHEET Selection
        if "sheet" in allKeys:
            self.sheet = cusTable['sheet']
        else:
            self.sheet = settings['tableAnalysis']['sheet']

        # Table Columns Selection:
        if "columns" in allKeys:
            self.columns = cusTable['columns']
        else:
            self.columns = settings['tableAnalysis']['columns']

        # System Prompts are Added (NOT REPLACED)!
        if "systemPrompts" in allKeys:
            self.systemPrompts = cusTable['systemPrompts'] + settings['tableAnalysis']['systemPrompts']
        else:
            self.systemPrompts = settings['tableAnalysis']['systemPrompts']
        # User Prompts are Added (NOT REPLACED)!
        if "userPrompts" in allKeys:
            self.userPrompts = cusTable['userPrompts'] + settings['tableAnalysis']['userPrompts']
        else:
            self.userPrompts = settings['tableAnalysis']['userPrompts']

        # GETTING ALL OF THE DEFAULT SETTING FOR BUBBLE CHART
        self.isBubbleChartRequired = settings['tableAnalysis']["bubbleChart"]["generate"]
        self.bubbleChartType = settings['tableAnalysis']["bubbleChart"]["type"]
        self.bubbleChartColumns = settings['tableAnalysis']["bubbleChart"][self.bubbleChartType]
        self.zoomIn = settings['tableAnalysis']["bubbleChart"]["zoomIn"]
        self.zoomRange = settings['tableAnalysis']["bubbleChart"]["zoomRange"]
        self.DIA_LARGEST_BUBLE = settings['tableAnalysis']["bubbleChart"]["DIA_LARGEST_BUBLE"]
        self.X_RANGE_MAX = settings['tableAnalysis']["bubbleChart"]["X_RANGE_MAX"]

        # OVERIDING THE DEFAULT SETTINGS
        if "bubbleChart" in allKeys:
            bubbleKeys = list(cusTable["bubbleChart"].keys())
            if "generate" in bubbleKeys:
                self.isBubbleChartRequired = cusTable["bubbleChart"]["generate"]
            if "type" in bubbleKeys:
                self.bubbleChartType = cusTable["bubbleChart"]["type"]
                self.bubbleChartColumns = settings['tableAnalysis']["bubbleChart"][self.bubbleChartType]
            if "zoomIn" in bubbleKeys:
                self.zoomIn = cusTable["bubbleChart"]["zoomIn"]
            if "zoomRange" in bubbleKeys:
                self.zoomRange = cusTable["bubbleChart"]["zoomRange"]
            if "DIA_LARGEST_BUBLE" in bubbleKeys:
                self.DIA_LARGEST_BUBLE = cusTable["bubbleChart"]["DIA_LARGEST_BUBLE"]
            if "X_RANGE_MAX" in bubbleKeys:
                self.X_RANGE_MAX = cusTable["bubbleChart"]["X_RANGE_MAX"]

        # If Bubble Chart Needs to be Generated
        if self.isBubbleChartRequired:
            # Chart Settings!
            self.FIG_HEIGTH = 500
            self.FIG_WIDTH = 500 + int((self.X_RANGE_MAX - 1) * 500)

            # Custom Legends !!!
            self.LEGEND_RADIUS = 0.008
            self.LEGEND_TOP_X = 1 + 0.04
            self.LEGEND_TOP_Y = 0.96
            self.LEGEND_SPACING = 0.034

            # ANNOTATION
            self.WIDTH_1_ALPHABET = 0.014
            self.HEIGHT_ALPHABET = 0.03

            # ANNOTATION LEGENDS
            self.ANNOTATION_LEGEND_BOT_X = 1 + 0.04
            self.ANNOTATION_LEGEND_BOT_Y = 0.01
            self.ANNOTATION_LEGEND_SPACING = 0.025

        self.settings = settings
        self.gptModel = settings["openAI"]["model"]
        self.apiPrompt = []
        self.chatGPTResponse = ""

        self.figures = []

    # Convert Table to CSV format for Chat-GPT
    def convertTableAnalysis(self, table):
        aTable = table.copy(deep=True)
        if self.rescale:
            for col in aTable.columns:
                if is_numeric_dtype(aTable[col]):
                    aTable[col] = ((aTable[col] - aTable[col].min()) / (aTable[col].max() - aTable[col].min()) *
                                   self.settings['tableAnalysis']['rescaleRange']).astype(int)

        # Adding The First Row Of TABLE!
        textTable = "\n \n Table: [ \n"
        for col in aTable.columns:
            textTable = textTable + col + ','
        textTable = textTable + " \n "

        for row in range(len(aTable.index)):
            for col in aTable.columns:
                textTable = textTable + str(aTable[col].iloc[row]) + ','
            textTable = textTable + " \n "

        textTable = textTable + "] \n"
        return textTable

    # Generating The Text Prompt for The Table!
    def generateTablePrompts(self, excelManager):
        table = self.requestingTable(excelManager, self.columns)
        txtTable = self.convertTableAnalysis(table)

        # Taking The Table Prompt First!
        systemPrompt = ""
        for prompt in self.systemPrompts:
            systemPrompt = systemPrompt + prompt + " \n "

        # Getting The Target Company Details!
        systemPrompt = systemPrompt + " \n Target Company: \n"
        for key in self.settings['target']:
            systemPrompt = systemPrompt + key + ": " + self.settings['target'][key] + "\n"

            # Getting The Report Generator Company!
        systemPrompt = systemPrompt + " \n Report Generated by Company: \n"
        for key in self.settings['generator']:
            systemPrompt = systemPrompt + key + ": " + self.settings['generator'][key] + "\n"

            # Getting The Details on Table Columns!
        systemPrompt = systemPrompt + " \n Detail on Table Columns: \n" + self.settings['tableInfo']['prompt'] + " \n "
        for col in self.columns:
            if col in list(self.settings['tableInfo'].keys()):
                systemPrompt = systemPrompt + col + ": " + self.settings['tableInfo'][col] + " \n "

        # Getting The Table Prompts
        systemPrompt = systemPrompt + " \n " + txtTable

        # Generating The User Prompt
        userPrompt = ""
        for line in self.userPrompts:
            userPrompt = userPrompt + line + "\n"

        apiMessage = [
            {"role": "system", "content": systemPrompt},
            {"role": "user", "content": userPrompt},
        ]
        self.apiPrompt = apiMessage
        return apiMessage

    # This function should return selected requested Columns of the Table!
    def requestingTable(self, excelManager, columns):
        table = excelManager.requestSheet(self.excel, self.sheet)
        # CHECK IF THE TABLE HAS ALL OF THE REQUIRED COLUMNS!!!
        for reqColumn in columns:
            found = False
            for col in table.columns:
                if col == reqColumn:
                    found = True
                    continue
            if not found:
                print("Some Columns Missing in the Table")
                print(f" This column:{reqColumn} is missing from the table")
                print("Program Terminating")
                sys.exit()

        # returning The Requested Table Columns!
        return table[columns]

    def analyzeGPT(self, client):

        if self.apiPrompt == []:
            print("The System Prompt for CHAT-GPT is Empty, Table ANALYSIS")
            print("CHECK analyzeGPT function of tableAnalysis")
            print("Program Terminating!!!")
            sys.exit()

        chat_completion = client.chat.completions.create(messages=self.apiPrompt, model=self.gptModel)
        self.chatGPTResponse = chat_completion.choices[0].message.content
        print(self.chatGPTResponse)
        print("1-Accept Response")
        print("2-Change Response")
        change_choice = input("Enter the number of your choice: ")
        if change_choice == "1":
            pass
        elif change_choice == "2":
            print("The current prompts are: ", self.apiPrompt)
            print('1- Make it brief')
            print('2- Add more details')
            print('3- Make it formal')
            print('4-Enter your feedback on the response')
            change_to_prompt = input("Enter the number of your choice: ")
            if change_to_prompt == "1":
                """ Make it brief"""
                self.apiPrompt[0]['content'] = ("This is the response for the analysis of the table. Please provide a "
                                               "brief response in one paragraph") + self.chatGPTResponse
                chat_completion = client.chat.completions.create(messages=self.apiPrompt, model=self.gptModel)
                self.chatGPTResponse = chat_completion.choices[0].message.content

                pass

            elif change_to_prompt == "2":

                self.apiPrompt[0]['content'] = ("This is the response for the analysis of the table. Please provide a "
                                                "detailed analysis.") + self.chatGPTResponse
                chat_completion = client.chat.completions.create(messages=self.apiPrompt, model=self.gptModel)
                self.chatGPTResponse = chat_completion.choices[0].message.content
            elif change_to_prompt == "3":
                self.apiPrompt[0]['content'] = ("This is the response for the analysis of the table. Please make it"
                                                "formal") + self.chatGPTResponse
                chat_completion = client.chat.completions.create(messages=self.apiPrompt, model=self.gptModel)
                self.chatGPTResponse = chat_completion.choices[0].message.content
            elif change_to_prompt == "4":

                feedback = input("Enter your feedback on the response: ")
                self.apiPrompt[0][
                    'content'] = 'The user gave this feedback:' + feedback + 'about this response:' + self.chatGPTResponse + 'please provide a new response based on the feedback and the previous response.'
                chat_completion = client.chat.completions.create(messages=self.apiPrompt, model=self.gptModel)
                self.chatGPTResponse = chat_completion.choices[0].message.content

            else:
                print("Invalid input")
                sys.exit()



        else:
            print("Invalid input")
            sys.exit()
        return self.chatGPTResponse

    def generateReport(self, document, excelManager):

        document.add_heading(self.title, 2)

        # Writing The Paragraphs
        splitParas = self.chatGPTResponse.split("\n")

        for para in splitParas:
            paragraph = document.add_paragraph(para)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if self.isBubbleChartRequired:
            fig = self.generateBubbleChart(excelManager, zoomVersion=False)
            #fig.show()
            #feedback = input("Enter your feedback on the bubble chart: ")
            self.figures.append(fig)


            # If zoomIn Figure Required!
            if self.zoomIn:
                self.figures.append(self.generateBubbleChart(excelManager, zoomVersion=True))

        imagePath = []
        for figInd in range(len(self.figures)):
            path = f'images/fig{figInd}.png'
            self.figures[figInd].write_image(path)
            imagePath.append(path)

        # Writing The Figure onto the Document
        for path in imagePath:
            document.add_picture(path, width=Inches(5.0))

    def getTitle(self):
        return self.title

    def getChatGPTResponse(self):
        return self.chatGPTResponse

    def getColorSchemeAndCodes(self, table):
        colorCode = []
        colors = []
        if self.bubbleChartType == "categoryBubbles":

            allCategories = sorted(list(set(table['Category'])))
            for category in table['Category']:
                for idx, refCat in enumerate(allCategories):
                    if category == refCat:
                        colorCode.append(idx)

            # Generating A Color Range!
            colorRange = []
            COLOR_CENTER_RANGE = 0.75
            colorStep = COLOR_CENTER_RANGE / len(allCategories)
            endGap = (1 - COLOR_CENTER_RANGE) / 2
            for i in range(len(allCategories)):
                colorRange.append(endGap + i * colorStep)
            colors = px.colors.sample_colorscale("turbo", colorRange)

        elif self.bubbleChartType == "rankBubbles":
            # Generating Color Catagories for TABLE!
            allCategories = ["Less than 0.4", "Between 0.4 and 0.5", "Between 0.5 and 0.6", "Greater than 0.6"]

            for idx in range(len(table.index)):
                rank = math.sqrt(
                    table['Toxicity Combined Score (Raw)'][idx] * table['Exposure Combined Score (Raw)'][idx])
                if rank < 0.4:
                    colorCode.append(0)
                elif rank < 0.5:
                    colorCode.append(1)
                elif rank < 0.6:
                    colorCode.append(2)
                else:
                    colorCode.append(3)

            # Generating The ColorRange!!!
            colors = ['rgb(147,184,99)', 'rgb(73,94,60)', 'rgb(215,78,9)', 'rgb(184,12,9)']
        else:
            print("You Have NOT Provided A valid TYPE FOR THE BUBBLE PLOT")

        return colorCode, colors, allCategories

    def generateBubbleChart(self, excelManager, zoomVersion=False):
        table = self.requestingTable(excelManager, self.bubbleChartColumns)

        # DROP NA VALUES!!!!
        table = table.dropna()

        colorCode, colors, allCategories = self.getColorSchemeAndCodes(table)
        table.insert(2, "colorCode", colorCode, True)

        # IF WE NEED ZOOM VERSION!!!
        if zoomVersion:
            L = self.zoomRange[0]
            H = self.zoomRange[1]
            table = table.loc[
                (table['Toxicity Combined Score (Raw)'] >= L) & (table['Toxicity Combined Score (Raw)'] <= H) & (
                        table['Exposure Combined Score (Raw)'] >= L) & (table['Exposure Combined Score (Raw)'] <= H)]

        # Set axes ranges
        columWidths = [1 / self.X_RANGE_MAX, (self.X_RANGE_MAX - 1) / self.X_RANGE_MAX]
        fig = make_subplots(rows=1, cols=2, column_widths=columWidths, horizontal_spacing=0)

        # Set axes ranges SUBPLOT 1
        fig.update_xaxes(range=[-0.005, 1], row=1, col=1)
        fig.update_yaxes(range=[-0.005, 1], row=1, col=1)
        # Set axes ranges SUBPLOT 2
        fig.update_xaxes(range=[1, self.X_RANGE_MAX], row=1, col=2)
        fig.update_yaxes(range=[-0.005, 1], row=1, col=2)

        if zoomVersion:
            fig.update_xaxes(range=[self.zoomRange[0] - 0.05, self.zoomRange[1] + 0.02], row=1, col=1)
            fig.update_yaxes(range=[self.zoomRange[0] - 0.05, self.zoomRange[1] + 0.02], row=1, col=1)

        # Records information of The Bubbles (ELIPS),
        # XCENTER, YCENTER, XRADIUS, YRADIUS
        bubbleArray = np.empty((0, 3), int)

        # Setting The Radius for the Zoomed Plot!!!
        LARGEST_BUBBLE_RADIUS = 0
        ALPHABET_HEIGHT = 0
        ALPHABET_WIDTH = 0
        if zoomVersion:
            LARGEST_BUBBLE_RADIUS = (self.zoomRange[1] - self.zoomRange[0]) * self.DIA_LARGEST_BUBLE
            ALPHABET_HEIGHT = (self.zoomRange[1] - self.zoomRange[0]) * self.HEIGHT_ALPHABET
            ALPHABET_WIDTH = (self.zoomRange[1] - self.zoomRange[0]) * self.WIDTH_1_ALPHABET
        else:
            LARGEST_BUBBLE_RADIUS = self.DIA_LARGEST_BUBLE
            ALPHABET_HEIGHT = self.HEIGHT_ALPHABET
            ALPHABET_WIDTH = self.WIDTH_1_ALPHABET

        # Data Bubbles !!!
        df = table.sort_values('Emerging Concern Score (raw)', ascending=False).reset_index(drop=True)
        ply_bubbles = {}
        for i in range(len(table['Chemical Name'])):
            xCenter = df['Toxicity Combined Score (Raw)'].iloc[i]
            yCenter = df['Exposure Combined Score (Raw)'].iloc[i]
            radius = df['Emerging Concern Score (raw)'].iloc[i] / df[
                'Emerging Concern Score (raw)'].max() * LARGEST_BUBBLE_RADIUS / 2
            bubbleArray = np.append(bubbleArray, [[xCenter, yCenter, radius]], axis=0)

            ply_bubbles['shape_' + str(i)] = go.layout.Shape(type="circle",
                                                             # xref="paper", yref="paper",
                                                             x0=xCenter - radius,
                                                             y0=yCenter - radius,
                                                             x1=xCenter + radius,
                                                             y1=yCenter + radius,
                                                             opacity=0.8,
                                                             layer="below",
                                                             line_width=5,
                                                             line_color=colors[df['colorCode'].iloc[i]],
                                                             fillcolor=colors[df['colorCode'].iloc[i]])

        # lst_shapes=list(ply_bubbles.values())
        # fig.update_layout(shapes=lst_shapes, width=FIG_WIDTH, height=FIG_HEIGTH,margin=dict(l=50, r=50, t=50, b=50) , plot_bgcolor='#F8F8F8')

        ply_legends = {}
        for i in range(len(allCategories)):
            xCenter = self.LEGEND_TOP_X
            yCenter = self.LEGEND_TOP_Y - i * self.LEGEND_SPACING
            # xRadius
            ply_legends['legend_' + str(i)] = go.layout.Shape(type="circle",
                                                              # xref="paper", yref="paper",
                                                              x0=xCenter - self.LEGEND_RADIUS,
                                                              y0=yCenter - self.LEGEND_RADIUS,
                                                              x1=xCenter + self.LEGEND_RADIUS,
                                                              y1=yCenter + self.LEGEND_RADIUS,
                                                              opacity=1,
                                                              layer="above",
                                                              line_width=4,
                                                              line_color=colors[i],
                                                              fillcolor=colors[i], xref='x2', yref='y2')

        # Custom Legend LABELS !!!
        xLabelPos = []
        yLabelPos = []
        for i in range(len(allCategories)):
            xLabelPos.append(self.LEGEND_TOP_X + 0.02)
            yLabelPos.append(self.LEGEND_TOP_Y - i * self.LEGEND_SPACING)
        fig.add_trace(go.Scatter(
            x=xLabelPos,
            y=yLabelPos,
            text=allCategories,
            mode="text",
            textfont_size=10,
            textposition="middle right"
        ), row=1, col=2)

        lst_shapes = list({**ply_bubbles, **ply_legends}.values())
        fig.update_layout(shapes=lst_shapes, width=self.FIG_WIDTH, height=self.FIG_HEIGTH,
                          margin=dict(l=50, r=50, t=50, b=50),
                          plot_bgcolor='#F8F8F8')

        # SUBPLOT 1
        fig.update_xaxes(showline=True, linewidth=1, linecolor='black', showgrid=False, ticks="outside",
                         tickson="boundaries", ticklen=4, tickwidth=2, row=1, col=1)
        fig.update_yaxes(showline=True, linewidth=1, linecolor='black', showgrid=False, ticks="outside",
                         tickson="boundaries", ticklen=4, tickwidth=2, row=1, col=1)

        # SUBPLOT 2
        fig.update_xaxes(showline=False, showgrid=False, showticklabels=False, row=1, col=2)
        fig.update_yaxes(showline=False, showgrid=False, showticklabels=False, row=1, col=2)

        nextAnnotation = "A"
        annLegends = []
        for buble in range(len(bubbleArray)):
            chemName = df['Chemical Name'].iloc[buble]
            optLoc = self.findOptimalLocAnnotation(bubbleArray=bubbleArray, bubbleNo=buble, textHeight=ALPHABET_HEIGHT,
                                                   textWidth=ALPHABET_WIDTH * len(chemName))
            if not optLoc.size == 0:
                fig.add_annotation(x=optLoc[0], y=optLoc[1], text=chemName, showarrow=False, arrowhead=1,
                                   align='center', font=dict(size=10), row=1, col=1)
            else:
                optLoc = self.findOptimalLocAnnotation(bubbleArray=bubbleArray, bubbleNo=buble,
                                                       textHeight=ALPHABET_HEIGHT,
                                                       textWidth=ALPHABET_WIDTH * len(nextAnnotation))
                if not optLoc.size == 0:
                    fig.add_annotation(x=optLoc[0], y=optLoc[1], text=nextAnnotation, showarrow=False, arrowhead=1,
                                       align='center', font=dict(size=10), row=1, col=1)
                else:
                    fig.add_annotation(x=bubbleArray[buble, 0], y=bubbleArray[buble, 1], text=nextAnnotation,
                                       showarrow=False, arrowhead=1, align='center', font=dict(size=10), row=1, col=1)

                annLegends.append([nextAnnotation, chemName])
                nextAnnotation = chr(ord(nextAnnotation) + 1)

        # ANNOTATION Legends LABELS HEADINGS !!!
        xANNPos = []
        yANNPos = []

        for i in range(len(annLegends)):
            xANNPos.append(self.ANNOTATION_LEGEND_BOT_X)
            yANNPos.append(self.ANNOTATION_LEGEND_BOT_Y + i * self.ANNOTATION_LEGEND_SPACING)

        # INVETING THE ANNOTATION POSITIONS
        xANNPos.reverse()
        yANNPos.reverse()

        fig.add_trace(go.Scatter(
            x=xANNPos,
            y=yANNPos,
            text=['<b>' + x[0] + '<b>' for x in annLegends],
            mode="text",
            textfont_size=10,
            textposition="middle left"), row=1, col=2)

        fig.add_trace(go.Scatter(
            x=np.array(xANNPos) + 0.015,
            y=yANNPos,
            text=['<b>' + ':' + '<b>' for i in range(len(annLegends))],
            mode="text",
            textfont_size=10,
            textposition="middle left"
        ), row=1, col=2)

        fig.add_trace(go.Scatter(
            x=np.array(xANNPos) + 0.030,
            y=yANNPos,
            text=[x[1] for x in annLegends],
            mode="text",
            textfont_size=9,
            textposition="middle right"), row=1, col=2)

        fig.add_vline(x=1, line_width=0.1, row=1, col=2)
        fig.update_layout(showlegend=False, xaxis_title="Relative Toxicity",
                          yaxis_title="Relative Likelihood of Exposure")

        # feedback = input("Enter your feedback on the bubble chart: ")

        # imagePath = "images/fig.png"
        # fig.write_image(imagePath)
        return fig

    @staticmethod
    def findOptimalLocAnnotation(bubbleArray, bubbleNo, textHeight, textWidth):

        xCent = bubbleArray[bubbleNo, 0]
        yCent = bubbleArray[bubbleNo, 1]
        radius = bubbleArray[bubbleNo, 2]
        xRange = np.linspace(xCent - radius, xCent + radius, num=30)
        yRange = np.linspace(yCent - radius, yCent + radius, num=30)
        xx, yy = np.meshgrid(xRange, yRange)
        xx = xx.flatten()
        yy = yy.flatten()

        # CHECK WHICH POINT CAN BE USED AS LABEL CENTER
        posInCircle = np.empty((0, 2))
        for i in range(len(xx)):

            # GET ALL OF THE POINTS i-e. TOP LEFT - TOP RIGHT - BOTTOM RIGHT - BOTTOM LEFT
            txtCorners = np.array([[xx[i] - textWidth / 2, yy[i] + textHeight / 2],
                                   [xx[i] + textWidth / 2, yy[i] + textHeight / 2],
                                   [xx[i] - textWidth / 2, yy[i] - textHeight / 2],
                                   [xx[i] + textWidth / 2, yy[i] - textHeight / 2]])

            # All of The Points MUST BE INSIDE THE CIRCLE !!!
            textFitsInside = True
            for corIdx in range(len(txtCorners)):
                # Check If inside the Circle!
                if (xCent - txtCorners[corIdx, 0]) ** 2 + (yCent - txtCorners[corIdx, 1]) ** 2 > radius ** 2:
                    textFitsInside = False
                    break
            if textFitsInside:
                posInCircle = np.append(posInCircle, [[xx[i], yy[i]]], axis=0)

        # Avoid Points Which are overLapped by Other Circles!!!
        validLocs = np.empty((0, 2))
        for i in range(len(posInCircle)):

            # GET ALL OF THE POINTS i-e. TOP LEFT - TOP RIGHT - BOTTOM RIGHT - BOTTOM LEFT
            txtCorners = np.array([[posInCircle[i, 0] - textWidth / 2, posInCircle[i, 1] + textHeight / 2],
                                   [posInCircle[i, 0] + textWidth / 2, posInCircle[i, 1] + textHeight / 2],
                                   [posInCircle[i, 0] - textWidth / 2, posInCircle[i, 1] - textHeight / 2],
                                   [posInCircle[i, 0] + textWidth / 2, posInCircle[i, 1] - textHeight / 2]])

            # All of The Points MUST BE OUSIDE THE OTHER CIRCLES!!!
            noBubbleOverlaps = True

            # Confirm That Other Points are not Overlapping!
            for bIdx in range(len(bubbleArray)):
                # Skip for the Same Bubble
                if bIdx == bubbleNo:
                    continue

                # Check if Text Box Corners are inside Other Bubbles
                for corIdx in range(len(txtCorners)):

                    if (bubbleArray[bIdx, 0] - txtCorners[corIdx, 0]) ** 2 + (
                            bubbleArray[bIdx, 1] - txtCorners[corIdx, 1]) ** 2 < bubbleArray[bIdx, 2] ** 2:
                        noBubbleOverlaps = False
                        break

                """
    
                # If X is in Range of the Text Annotation
                if bubbleArray[bIdx, 0] > txtCorners[0,0] and bubbleArray[bIdx, 0] < txtCorners[0,1]:   # If Bubble X-Center Falls between L & R Edge of TXT
                    # if the Bubble Crosses into the Top or bottom of the TXT BOX ie. distance is less than radius
                    if abs(bubbleArray[bIdx, 1] - txtCorners[0,1]) < bubbleArray[bIdx, 2] or abs(bubbleArray[bIdx, 1] - txtCorners[3,1]) < bubbleArray[bIdx, 2]:
                        noBubbleOverlaps = False
                """

                # If One Bubble Overlaps No need to check for other Locations
                if not noBubbleOverlaps:
                    break

            # Append the location if Other Bubbles are not Overlapping!!!
            if noBubbleOverlaps:
                validLocs = np.append(validLocs, [[posInCircle[i, 0], posInCircle[i, 1]]], axis=0)
        # return validLocs
        # Choose The location Nearest to the center!!!
        if len(validLocs) > 0:
            minIdx = np.argmin((validLocs[:, 0] - xCent) ** 2 + (validLocs[:, 1] - yCent) ** 2)
            return validLocs[minIdx, :]
        else:
            return np.empty((0, 2))
